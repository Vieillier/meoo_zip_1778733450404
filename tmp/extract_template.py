from docx import Document
import sys

doc = Document('/home/project/assets/审核意见通过书-模板.docx')

print("=== 文档段落内容 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n=== 文档表格内容 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n表格 {t_idx}:")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(f"  行: {row_text}")
